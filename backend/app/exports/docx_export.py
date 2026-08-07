"""Customer-pack styled DOCX negotiation letter."""

from __future__ import annotations

import os
from collections.abc import Iterable, Mapping, Sequence
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.exports.common import money_label, percentage_label
from app.exports.letter_context import (
    LetterFacts,
    build_letter_facts,
    deterministic_line_sentence,
)

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = RGBColor(23, 50, 77)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(122, 90, 0)
GRAY = RGBColor(91, 101, 111)
WHITE = RGBColor(255, 255, 255)


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    current = tc_pr.find(qn("w:tcMar"))
    if current is not None:
        tc_pr.remove(current)
    margins = OxmlElement("w:tcMar")
    for edge, value in CELL_MARGINS_DXA.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tc_pr.append(margins)


def _replace_child(parent: Any, tag: str, child: Any) -> None:
    current = parent.find(qn(tag))
    if current is not None:
        parent.remove(current)
    parent.append(child)


def _set_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"DOCX table widths must sum to {PAGE_WIDTH_DXA} DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    _replace_child(tbl_pr, "w:tblW", tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    _replace_child(tbl_pr, "w:tblInd", tbl_ind)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _replace_child(tbl_pr, "w:tblLayout", layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _replace_child(tc_pr, "w:tcW", tc_w)


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_page_field(paragraph: Any) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 11, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _header_footer(document: Document, facts: LetterFacts) -> None:
    section = document.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_geometry(header_table, (4680, 4680))
    header_table.style = "Table Grid"
    for cell in header_table.rows[0].cells:
        _set_cell_shading(cell, "FFFFFF")
    left = header_table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(left.add_run("ClaimGuard"), size=9, color=NAVY, bold=True)
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(right.add_run(f"Case {facts.case_reference}"), size=9, color=GRAY)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        paragraph.add_run(f"ClaimGuard | {facts.report_date} | "),
        size=8,
        color=GRAY,
    )
    _add_page_field(paragraph)
    for run in paragraph.runs:
        _set_run_font(run, size=8, color=GRAY)


def _add_key_value_table(
    document: Document,
    rows: Iterable[tuple[str, str]],
    *,
    widths: tuple[int, int] = (1701, 7659),
    header: tuple[str, str] | None = None,
) -> Any:
    values = list(rows)
    table = document.add_table(rows=1 if header else 0, cols=2)
    table.style = "Table Grid"
    if header:
        for idx, value in enumerate(header):
            cell = table.rows[0].cells[idx]
            _set_cell_shading(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            _set_run_font(paragraph.add_run(value), size=9, color=NAVY, bold=True)
        _set_repeat_table_header(table.rows[0])
    for label, value in values:
        cells = table.add_row().cells
        _set_cell_shading(cells[0], "F4F6F9")
        _set_run_font(cells[0].paragraphs[0].add_run(label), size=9, color=NAVY, bold=True)
        _set_run_font(cells[1].paragraphs[0].add_run(value), size=9, color=RGBColor(31, 41, 51))
    _set_table_geometry(table, widths)
    return table


def _add_first_page(document: Document, facts: LetterFacts) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(0)
    _set_run_font(kicker.add_run("CLAIM NEGOTIATION BRIEF"), size=9, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    _set_run_font(
        title.add_run(f"Challenge Price: {money_label(facts.summary.challenge_price_net)} net"),
        size=28,
        color=NAVY,
        bold=True,
    )

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    _set_run_font(
        subtitle.add_run(
            "The proposed payable amount after deterministic review of the approved invoice lines."
        ),
        size=13.5,
        color=GRAY,
    )

    _add_key_value_table(
        document,
        (
            ("Case", facts.case_reference),
            ("Claim", facts.claim_number),
            ("Invoice", facts.invoice_references),
            ("Vehicle", facts.vehicle_registration),
            ("Prepared for", facts.recipient),
            ("Prepared by", facts.paying_insurer),
            ("Liability", f"{facts.liability_status} - human confirmed"),
            ("Date", facts.report_date),
        ),
    )


def _add_letter_body(document: Document, facts: LetterFacts) -> None:
    document.add_heading("Position and proposal", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Dear Claims Team,")

    paragraph = document.add_paragraph()
    paragraph.add_run(
        f"We have reviewed invoice {facts.invoice_references} for vehicle "
        f"{facts.vehicle_registration}. The recorded liability position is "
        f"{facts.liability_status}; it was confirmed by {facts.liability_confirmed_by}"
        f"{f' on {facts.liability_confirmed_at}' if facts.liability_confirmed_at else ''}. "
        "This letter reports that human-confirmed position and does not infer liability from the invoice."
    )

    paragraph = document.add_paragraph()
    paragraph.add_run(
        f"The Invoice Price is {money_label(facts.summary.invoice_price_net)} net. "
        f"Our Challenge Price is {money_label(facts.summary.challenge_price_net)} net "
        f"(the proposed payable amount), producing a Challenge Amount of "
        f"{money_label(facts.summary.challenge_amount_net)} net."
    )

    document.add_heading("Financial summary", level=1)
    _add_key_value_table(
        document,
        (
            ("Invoice Price", f"{money_label(facts.summary.invoice_price_net)} net"),
            (
                "Challenge Price",
                f"{money_label(facts.summary.challenge_price_net)} net - proposed payable",
            ),
            ("Challenge Amount", f"{money_label(facts.summary.challenge_amount_net)} net"),
            ("VAT impact", money_label(facts.summary.vat_impact)),
            ("Gross effect", money_label(facts.summary.gross_effect)),
            ("MOT / non-VAT", f"{money_label(facts.summary.mot_non_vat_total)} outside VAT"),
            ("Reduction", percentage_label(facts.summary.challenge_percentage)),
        ),
        widths=(1701, 7659),
    )
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(6)
    _set_run_font(
        note.add_run(
            "All line evidence and the Challenge Price are stated net. VAT impact is shown separately; "
            "MOT and other non-VAT charges remain outside the VAT computation."
        ),
        size=9,
        color=GRAY,
        italic=True,
    )

    document.add_heading("Approved challenged lines", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("Item", "Invoice net", "Challenge Price", "Challenge Amount", "Basis")
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        _set_run_font(paragraph.add_run(label), size=8.5, color=NAVY, bold=True)
    _set_repeat_table_header(table.rows[0])
    for line in facts.lines:
        cells = table.add_row().cells
        values = (
            line.description,
            money_label(line.invoice_net),
            money_label(line.challenge_price_net),
            money_label(line.challenge_amount_net),
            line.benchmark_source,
        )
        for idx, value in enumerate(values):
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if idx in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_run_font(paragraph.add_run(value), size=8.3, color=RGBColor(31, 41, 51))
    _set_table_geometry(table, (3168, 1152, 1440, 1440, 2160))

    document.add_heading("Evidence by line", level=1)
    for index, line in enumerate(facts.lines, start=1):
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(4 if index > 1 else 0)
        heading.paragraph_format.space_after = Pt(2)
        heading.paragraph_format.keep_with_next = True
        _set_run_font(
            heading.add_run(f"{index}. {line.description}"),
            size=10.5,
            color=DARK_BLUE,
            bold=True,
        )
        paragraph = document.add_paragraph(deterministic_line_sentence(line))
        paragraph.paragraph_format.space_after = Pt(5)

    document.add_heading("Response requested", level=1)
    response = document.add_paragraph()
    response.add_run(
        "Please confirm acceptance of the Challenge Price or provide line-specific evidence supporting "
        "the disputed charges. The calculation can be reconciled directly to the approved rows above."
    )
    response.add_run("\nYours faithfully - ")
    _set_run_font(response.add_run(facts.paying_insurer), bold=True, color=NAVY)
    response.add_run(" claims team.")


def build_negotiation_docx(result: Mapping[str, Any], output_path: str | Path) -> Path:
    """Build the fact-only negotiation letter using the resolved design tokens."""

    facts = build_letter_facts(result)
    document = Document()
    _style_document(document)
    _header_footer(document, facts)
    _add_first_page(document, facts)
    _add_letter_body(document, facts)

    document.core_properties.title = f"ClaimGuard negotiation letter - {facts.case_reference}"
    document.core_properties.subject = "Motor invoice Challenge Price and supporting evidence"
    document.core_properties.author = "ClaimGuard"
    document.core_properties.last_modified_by = "ClaimGuard"
    document.core_properties.keywords = "ClaimGuard, challenge price, motor claim, negotiation"

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with NamedTemporaryFile(
        prefix=f".{path.stem}-", suffix=".docx", dir=path.parent, delete=False
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)
    return path
