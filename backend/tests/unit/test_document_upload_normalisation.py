import io
from pathlib import Path
from types import SimpleNamespace

import fitz
import pytest
from docx import Document as DocxDocument

import app.services.document_processing as document_processing


def _build_invoice_docx() -> bytes:
    """Build a small in-memory .docx resembling a repair invoice."""

    document = DocxDocument()
    document.add_heading("INVOICE 12345", level=1)
    document.add_paragraph("Supplier: Acme Bodyshop Ltd")
    document.add_paragraph("Customer: Jane Doe")

    table = document.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Description"
    header[1].text = "Qty"
    header[2].text = "Price"
    line_items = [
        ("Front bumper replacement", "1", "250.00"),
        ("Headlamp assembly, offside", "1", "180.00"),
        ("Paint and refinish - bumper", "2", "95.00"),
    ]
    for description, qty, price in line_items:
        row = table.add_row().cells
        row[0].text = description
        row[1].text = qty
        row[2].text = price

    document.add_paragraph("Subtotal: 620.00")
    document.add_paragraph("VAT: 124.00")
    document.add_paragraph("Total: 744.00")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_is_converted_to_pdf_before_storage(
    tmp_path: Path, monkeypatch
) -> None:
    monkeypatch.setattr(document_processing.shutil, "which", lambda _: "/usr/bin/soffice")

    def fake_run(command, **kwargs):
        output_dir = Path(command[command.index("--outdir") + 1])
        document = fitz.open()
        document.new_page().insert_text((72, 72), "Invoice INV-1")
        document.save(output_dir / "source.pdf")
        document.close()
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(document_processing.subprocess, "run", fake_run)

    normalised = document_processing.normalise_document_upload(
        "repair.docx", b"PK\x03\x04fake-office-package"
    )

    assert normalised.content.startswith(b"%PDF-")
    assert normalised.stored_filename == "repair.pdf"
    assert normalised.source_format == "docx-libreoffice"


def test_docx_falls_back_to_python_pdf_without_libreoffice(monkeypatch) -> None:
    monkeypatch.setattr(document_processing.shutil, "which", lambda _: None)

    docx_bytes = _build_invoice_docx()
    normalised = document_processing.normalise_document_upload("invoice-12345.docx", docx_bytes)

    assert normalised.content.startswith(b"%PDF-")
    assert normalised.stored_filename == "invoice-12345.pdf"
    assert normalised.source_format == "docx-python"

    extracted = fitz.open(stream=normalised.content, filetype="pdf")
    try:
        text = "\n".join(page.get_text() for page in extracted)
    finally:
        extracted.close()

    assert "INVOICE 12345" in text
    assert "Front bumper replacement" in text


def test_doc_without_libreoffice_raises_clear_error(monkeypatch) -> None:
    monkeypatch.setattr(document_processing.shutil, "which", lambda _: None)

    with pytest.raises(ValueError, match="DOC \\(legacy\\)"):
        document_processing.normalise_document_upload(
            "repair.doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1fake-legacy-doc"
        )


def test_pdf_upload_is_not_rewritten() -> None:
    content = b"%PDF-1.7\nexample"
    normalised = document_processing.normalise_document_upload("repair.pdf", content)
    assert normalised.content == content
    assert normalised.stored_filename == "repair.pdf"
    assert normalised.source_format == "pdf"
