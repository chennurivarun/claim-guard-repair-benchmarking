from __future__ import annotations

import json
import shutil
import sqlite3
from copy import deepcopy
from datetime import UTC, datetime
from decimal import Decimal

import pdfplumber
import pytest
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook

from app.exports import (
    ExportValidationError,
    backup_sqlite,
    build_case_workbook,
    build_json_bytes,
    build_negotiation_docx,
    build_negotiation_pdf,
    compute_financial_summary,
)


@pytest.fixture()
def case_result() -> dict:
    return {
        "report_date": "2026-07-17",
        "case": {
            "case_reference": "CG-2026-0001",
            "claim_number": "CLM-7788",
            "paying_insurer_name": "Northstar Insurance",
            "claiming_insurer_name": "Harbour Mutual",
            "accident_date": "2026-05-12",
        },
        "liability": {
            "status": "ADMITTED",
            "human_confirmed": True,
            "confirmed_by": "A. Handler",
            "confirmed_at": "2026-07-16T15:30:00Z",
            "evidence": [
                {
                    "id": "ev-1",
                    "type": "statement",
                    "source": "handler",
                    "summary": "Liability admitted in the claim file.",
                }
            ],
        },
        "invoices": [
            {
                "id": "inv-1",
                "invoice_number": "INV-91283",
                "invoice_date": "2026-06-01",
                "repairer": "St Albans Garage",
                "vehicle_registration": "AB12 CDE",
                "net_total": "300.00",
                "vat_total": "40.00",
                "non_vat_total": "100.00",
                "gross_total": "340.00",
            }
        ],
        "pages": [
            {
                "document_id": "doc-1",
                "page_number": 1,
                "classification": "INVOICE",
                "confidence": "0.99",
            }
        ],
        "lines": [
            {
                "id": "line-1",
                "invoice_id": "inv-1",
                "description": "Front brake pads",
                "invoice_net": "100.00",
                "vat_rate": "0.20",
                "kind": "PART",
            },
            {
                "id": "line-2",
                "invoice_id": "inv-1",
                "description": "Wheel alignment under benchmark",
                "invoice_net": "40.00",
                "vat_rate": "0.20",
                "kind": "LABOUR",
            },
            {
                "id": "line-3",
                "invoice_id": "inv-1",
                "description": "Unapproved replacement engine",
                "invoice_net": "60.00",
                "vat_rate": "0.20",
                "kind": "PART",
            },
            {
                "id": "line-4",
                "invoice_id": "inv-1",
                "description": "MOT test",
                "invoice_net": "100.00",
                "vat_rate": "0.20",
                "kind": "MOT",
                "is_mot": True,
            },
        ],
        "checks": [
            {"id": "check-1", "check_type": "invoice_total", "status": "PASS", "severity": "INFO"}
        ],
        "mappings": [
            {
                "line_id": "line-1",
                "status": "MATCH",
                "ontology_item_id": "PART-045",
                "mapping_confidence": "0.94",
            }
        ],
        "comparisons": [
            {
                "line_id": "line-1",
                "ontology_price_net": "50.00",
                "historical_median_net": "55.00",
                "historical_count": 8,
            }
        ],
        "challenges": [
            {
                "line_id": "line-1",
                "approved": True,
                "challengeable": True,
                "challenge_price_net": "55.00",
                "challenge_amount_net": "999.00",
                "ontology_price_net": "50.00",
                "historical_median_net": "55.00",
                "historical_count": 8,
                "benchmark_source": "conservative supported evidence",
                "reason": "Approved ontology and settled comparables support a lower price.",
            },
            {
                "line_id": "line-2",
                "approved": True,
                "challengeable": True,
                "challenge_price_net": "50.00",
                "challenge_amount_net": "12.00",
                "benchmark_source": "approved ontology",
            },
            {
                "line_id": "line-3",
                "approved": False,
                "challengeable": True,
                "challenge_price_net": "40.00",
                "challenge_amount_net": "20.00",
                "benchmark_source": "provisional research",
            },
            {
                "line_id": "line-4",
                "approved": True,
                "challengeable": True,
                "challenge_price_net": "90.00",
                "challenge_amount_net": "10.00",
                "benchmark_source": "approved MOT schedule",
            },
        ],
        "evidence": [
            {
                "id": "price-1",
                "type": "historical",
                "source": "settled repair observation",
                "captured_at": "2026-07-01",
            }
        ],
        "versions": [
            {"type": "policy", "version": "challenge-policy-1.4", "status": "ACTIVE"},
            {"type": "ontology", "version": "ontology-3", "status": "PUBLISHED"},
        ],
        "audit": [
            {
                "id": "audit-1",
                "timestamp": "2026-07-16T15:30:00Z",
                "actor": "A. Handler",
                "action": "APPROVE",
            }
        ],
    }


def test_financial_summary_recomputes_positive_approved_lines(case_result: dict) -> None:
    summary = compute_financial_summary(case_result)
    assert summary.invoice_price_net == Decimal("300.00")
    assert summary.challenge_amount_net == Decimal("55.00")
    assert summary.challenge_price_net == Decimal("245.00")
    assert summary.vat_impact == Decimal("9.00")
    assert summary.gross_effect == Decimal("64.00")
    assert summary.proposed_payable_gross == Decimal("276.00")
    assert summary.challenge_percentage == Decimal("18.33")
    assert summary.challenged_line_count == 2


def test_json_export_is_stable_and_includes_computed_facts(case_result: dict) -> None:
    first = build_json_bytes(case_result)
    second = build_json_bytes(case_result)
    assert first == second
    payload = json.loads(first)
    assert payload["case"]["case_reference"] == "CG-2026-0001"
    assert payload["computed_export_summary"]["challenge_price_net"] == "245.00"
    assert payload["computed_export_summary"]["vat_impact"] == "9.00"
    assert [line["line_id"] for line in payload["computed_letter_lines"]] == [
        "line-1",
        "line-4",
    ]
    assert payload["computed_letter_lines"][0]["challenge_amount_net"] == "45.00"


def test_xlsx_contains_all_audit_sheets_and_typed_computed_values(
    case_result: dict, tmp_path
) -> None:
    path = build_case_workbook(case_result, tmp_path / "case.xlsx")
    workbook = load_workbook(path, data_only=False)
    assert workbook.sheetnames == [
        "Claim",
        "Liability",
        "Invoices",
        "Pages",
        "Lines",
        "Checks",
        "Mappings",
        "Comparisons",
        "Challenges",
        "Evidence",
        "Versions",
        "Audit",
    ]
    assert all(workbook[name].freeze_panes == "A5" for name in workbook.sheetnames)

    claim = workbook["Claim"]
    claim_values = {
        claim.cell(row, 2).value: claim.cell(row, 3).value for row in range(5, claim.max_row + 1)
    }
    assert claim_values["challenge_price_net"] == 245
    assert claim_values["challenge_amount_net"] == 55

    challenges = workbook["Challenges"]
    headers = {
        challenges.cell(4, column).value: column for column in range(1, challenges.max_column + 1)
    }
    rows = {
        challenges.cell(row, headers["line_id"]).value: row
        for row in range(5, challenges.max_row + 1)
    }
    assert challenges.cell(rows["line-1"], headers["included_in_letter"]).value is True
    assert challenges.cell(rows["line-1"], headers["computed_challenge_amount_net"]).value == 45
    assert challenges.cell(rows["line-4"], headers["computed_vat_impact"]).value == 0
    assert challenges.cell(rows["line-3"], headers["included_in_letter"]).value is False

    comparisons = workbook["Comparisons"]
    comparison_headers = {
        comparisons.cell(4, column).value: column for column in range(1, comparisons.max_column + 1)
    }
    comparison_rows = {
        comparisons.cell(row, comparison_headers["line_id"]).value: row
        for row in range(5, comparisons.max_row + 1)
    }
    line_one = comparison_rows["line-1"]
    assert comparisons.cell(line_one, comparison_headers["invoice_net"]).value == 100
    assert comparisons.cell(line_one, comparison_headers["challenge_price_net"]).value == 55
    assert comparisons.cell(line_one, comparison_headers["challenge_amount_net"]).value == 45


def test_xlsx_normalises_timezone_aware_datetimes(case_result: dict, tmp_path) -> None:
    case_result["audit"][0]["timestamp"] = datetime(2026, 7, 17, 9, 30, tzinfo=UTC)

    path = build_case_workbook(case_result, tmp_path / "aware.xlsx")

    workbook = load_workbook(path, data_only=True)
    assert workbook["Audit"].cell(5, 2).value is not None


def _docx_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_docx_letter_uses_fixed_terms_and_only_approved_lines(case_result: dict, tmp_path) -> None:
    path = build_negotiation_docx(case_result, tmp_path / "letter.docx")
    document = Document(path)
    text = _docx_text(document)
    assert "Challenge Price: £245.00 net" in text
    assert "proposed payable amount" in text
    assert "Challenge Amount of £55.00 net" in text
    assert "VAT impact" in text and "£9.00" in text
    assert "Gross effect" in text and "£64.00" in text
    assert "MOT and other non-VAT charges remain outside the VAT computation" in text
    assert "ADMITTED - human confirmed" in text
    assert "Front brake pads" in text
    assert "MOT test" in text
    assert "Wheel alignment under benchmark" not in text
    assert "Unapproved replacement engine" not in text

    for table in document.tables:
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        assert table_width is not None
        assert table_width.get(qn("w:w")) == "9360"
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360


def test_negotiation_letter_is_blocked_when_liability_is_denied(
    case_result: dict, tmp_path
) -> None:
    case_result["liability"]["status"] = "DENIED"
    case_result["liability"]["human_confirmed"] = True

    with pytest.raises(ExportValidationError, match="ADMITTED or SPLIT LIABILITY"):
        build_negotiation_docx(case_result, tmp_path / "denied.docx")


def test_reportlab_pdf_contains_same_fact_set(case_result: dict, tmp_path) -> None:
    path = build_negotiation_pdf(
        case_result,
        tmp_path / "letter.pdf",
        method="reportlab",
    )
    with pdfplumber.open(path) as document:
        text = "\n".join(page.extract_text() or "" for page in document.pages)
    compact_text = " ".join(text.split())
    assert "Challenge Price: £245.00 net" in text
    assert "proposed payable amount" in text
    assert "VAT impact" in text
    assert "MOT and other non-VAT charges remain outside the VAT computation" in compact_text
    assert "Front brake pads" in text
    assert "MOT test" in text
    assert "Wheel alignment under benchmark" not in text
    assert "Unapproved replacement engine" not in text


def test_libreoffice_pdf_conversion_when_available(case_result: dict, tmp_path) -> None:
    if not (shutil.which("soffice") or shutil.which("libreoffice")):
        pytest.skip("LibreOffice is not installed")
    path = build_negotiation_pdf(
        case_result,
        tmp_path / "letter-from-docx.pdf",
        method="libreoffice",
    )
    with pdfplumber.open(path) as document:
        text = " ".join(" ".join(page.extract_text().split()) for page in document.pages)
    assert "Challenge Price: £245.00 net" in text
    assert "Front brake pads" in text
    assert "Unapproved replacement engine" not in text


@pytest.mark.parametrize(
    "builder,extension", [(build_negotiation_docx, ".docx"), (build_negotiation_pdf, ".pdf")]
)
def test_letter_rejects_unconfirmed_liability(
    case_result: dict, tmp_path, builder, extension
) -> None:
    unconfirmed = deepcopy(case_result)
    unconfirmed["liability"]["human_confirmed"] = False
    with pytest.raises(ExportValidationError, match="human-confirmed liability"):
        builder(unconfirmed, tmp_path / f"blocked{extension}")


def test_sqlite_backup_uses_online_backup_and_passes_integrity_check(tmp_path) -> None:
    source = tmp_path / "source.db"
    with sqlite3.connect(source) as connection:
        connection.execute("PRAGMA journal_mode=WAL")
        connection.execute("CREATE TABLE audit (id INTEGER PRIMARY KEY, action TEXT NOT NULL)")
        connection.execute("INSERT INTO audit(action) VALUES (?)", ("APPROVE",))
        connection.commit()

    destination = backup_sqlite(source, tmp_path / "backup.db")
    with sqlite3.connect(destination) as connection:
        assert connection.execute("PRAGMA integrity_check").fetchone() == ("ok",)
        assert connection.execute("SELECT action FROM audit").fetchone() == ("APPROVE",)
